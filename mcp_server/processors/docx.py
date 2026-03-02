"""DOCX processor — text extraction with heading structure and table support."""

from __future__ import annotations

import io
import logging
from typing import List

from docx import Document as DocxDocument

logger = logging.getLogger("mcp_server.processors.docx")


def extract_docx_text(file_content: bytes) -> str:
    """Extract structured text from a DOCX document."""
    try:
        doc = DocxDocument(io.BytesIO(file_content))
        parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name.lower() if para.style else ""
            if "heading" in style:
                level = 1
                for ch in style:
                    if ch.isdigit():
                        level = int(ch)
                        break
                parts.append(f"\n{'#' * level} {text}\n")
            else:
                parts.append(text)

        # Extract tables
        for idx, table in enumerate(doc.tables):
            rows: List[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append(f"\n--- Table {idx + 1} ---")
                parts.extend(rows)

        return "\n".join(parts)

    except Exception as exc:
        logger.error(f"DOCX extraction failed: {exc}")
        return ""
